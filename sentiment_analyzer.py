import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import requests
from huggingface_hub import InferenceClient
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics import confusion_matrix, classification_report
import time
import base64
import io
from docx import Document
from docx.shared import Inches
from xhtml2pdf import pisa
import chardet
import re
from collections import Counter
import string
import subprocess
import sys
import seaborn as sns

# -------------------------
# Install NLTK if not available
# -------------------------
try:
    import nltk
    from nltk.corpus import stopwords
    # Download NLTK resources (if not already downloaded)
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt')

    try:
        nltk.data.find('corpora/stopwords')
    except LookupError:
        nltk.download('stopwords')

    import nltk
    nltk.download('punkt')
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize

except ImportError:
    st.warning("NLTK not found. Installing it now...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "nltk"])
    import nltk
    nltk.download('punkt')
    nltk.download('stopwords')
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize

# -------------------------
# Streamlit Page Config
# -------------------------
st.set_page_config(
    page_title="Sentiment Analyzer", 
    layout="wide",
    page_icon="📊"
)

# -------------------------
# Pre-configured Hugging Face API
# -------------------------
HF_API_KEY = "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"

# -------------------------
# Initialize session state variables
# -------------------------
if 'df_results' not in st.session_state:
    st.session_state.df_results = None
    
if 'comparison_results' not in st.session_state:
    st.session_state.comparison_results = None

# -------------------------
# Custom CSS Styling
# -------------------------
def inject_custom_css():
    st.markdown("""
    <style>
    /* Main styling */
    .main {
        background-color: #f8f9fa;
    }
    
    /* Header styling */
    .header {
        background: linear-gradient(135deg, #6e8efb, #a777e3);
        color: white;
        padding: 2rem;
        border-radius: 0.5rem;
        margin-bottom: 2rem;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    
    .header h1 {
        font-size: 2.5rem;
        font-weight: 700;
        margin-bottom: 0.5rem;
    }
    
    .header h2 {
        font-size: 1.5rem;
        font-weight: 300;
        opacity: 0.9;
    }
    
    /* Card styling */
    .card {
        background-color: white;
        border-radius: 0.5rem;
        padding: 1.5rem;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
        margin-bottom: 1.5rem;
        border-left: 4px solid #6e8efb;
    }
    
    /* Button styling */
    .stButton button {
        background: linear-gradient(135deg, #6e8efb, #a777e3);
        color: white;
        border: none;
        border-radius: 0.5rem;
        padding: 0.5rem 1rem;
        font-weight: 500;
        transition: all 0.3s ease;
    }
    
    .stButton button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
    }
    
    /* Sidebar styling */
    .css-1d391kg {
        background-color: #f8f9fa;
    }
    
    /* Progress bar styling */
    .stProgress > div > div {
        background: linear-gradient(90deg, #6e8efb, #a777e3);
    }
    
    /* Sentiment colors */
    .positive {
        color: #28a745;
        font-weight: 600;
    }
    
    .negative {
        color: #dc3545;
        font-weight: 600;
    }
    
    .neutral {
        color: #ffc107;
        font-weight: 600;
    }
    
    /* Tab styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    
    .stTabs [data-baseweb="tab"] {
        background-color: #f8f9fa;
        border-radius: 4px 4px 0 0;
        padding: 10px 16px;
        font-weight: 500;
    }
    
    .stTabs [aria-selected="true"] {
        background-color: white;
        border-bottom: 3px solid #6e8efb;
    }
    
    /* Text area styling */
    .stTextArea textarea {
        border-radius: 0.5rem;
        border: 1px solid #e9ecef;
        padding: 1rem;
    }
    
    /* Success message */
    .success-msg {
        background-color: #d4edda;
        color: #155724;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #28a745;
        margin-bottom: 1rem;
    }
    
    /* Info message */
    .info-msg {
        background-color: #d1ecf1;
        color: #0c5460;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #17a2b8;
        margin-bottom: 1rem;
    }
    </style>
    """, unsafe_allow_html=True)

# Inject custom CSS
inject_custom_css()

# -------------------------
# Improved CSV Reading with Encoding Detection
# -------------------------
def read_csv_with_encoding(uploaded_file, sample_size=1024):
    """Read CSV file with automatic encoding detection"""
    try:
        # First try reading with default utf-8
        return pd.read_csv(uploaded_file)
    except UnicodeDecodeError:
        try:
            # If utf-8 fails, try to detect encoding
            uploaded_file.seek(0)
            raw_data = uploaded_file.read(sample_size)
            uploaded_file.seek(0)
            
            # Detect encoding
            result = chardet.detect(raw_data)
            encoding = result['encoding']
            confidence = result['confidence']
            
            st.info(f"Detected encoding: {encoding} (confidence: {confidence:.2f})")
            
            # Try reading with detected encoding
            if encoding:
                return pd.read_csv(uploaded_file, encoding=encoding)
            else:
                # Fallback to latin-1 which handles most cases
                uploaded_file.seek(0)
                return pd.read_csv(uploaded_file, encoding='latin-1')
                
        except Exception as e:
            st.error(f"Failed to read CSV file: {str(e)}")
            return None
    except Exception as e:
        st.error(f"Error reading CSV file: {str(e)}")
        return None

# -------------------------
# Export Functions
# -------------------------
def create_html_export(df, title="Sentiment Analysis Results"):
    """Create HTML export of results"""
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>{title}</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; }}
            h1 {{ color: #2e86c1; }}
            table {{ border-collapse: collapse; width: 100%; }}
            th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
            th {{ background-color: #f2f2f2; }}
            tr:nth-child(even) {{ background-color: #f9f9f9; }}
            .positive {{ color: green; }}
            .negative {{ color: red; }}
            .neutral {{ color: orange; }}
        </style>
    </head>
    <body>
        <h1>{title}</h1>
        <p>Generated on: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        {df.to_html(classes='dataframe', escape=False, index=False)}
    </body>
    </html>
    """
    return html

def create_word_export(df, title="Sentiment Analysis Results"):
    """Create Word document export of results"""
    doc = Document()
    doc.add_heading(title, 0)
    
    # Add generation date
    doc.add_paragraph(f"Generated on: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Add table
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    
    # Add header row
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = str(column)
    
    # Add data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf_export(df, title="Sentiment Analysis Results"):
    """Create PDF export of results"""
    html = create_html_export(df, title)
    
    # Create PDF from HTML
    buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(html, dest=buffer)
    
    if pisa_status.err:
        st.error("PDF generation failed")
        return None
        
    buffer.seek(0)
    return buffer

def create_excel_export(df):
    """Create Excel export of results"""
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name='Sentiment Analysis', index=False)
        
        # Get workbook and worksheet
        workbook = writer.book
        worksheet = writer.sheets['Sentiment Analysis']
        
        # Add some formatting
        header_format = workbook.add_format({
            'bold': True,
            'text_wrap': True,
            'valign': 'top',
            'fg_color': '#D7E4BC',
            'border': 1
        })
        
        # Write the column headers with the defined format
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)
            
        # Auto-adjust columns' width
        for idx, col in enumerate(df.columns):
            series = df[col]
            max_len = max((
                series.astype(str).map(len).max(),
                len(str(col))
            )) + 2
            worksheet.set_column(idx, idx, max_len)
    
    buffer.seek(0)
    return buffer

def create_json_export(df):
    """Create JSON export of results"""
    import json
    return json.dumps(df.to_dict(orient='records'), indent=2)

# -------------------------
# Hugging Face Helpers with Pre-configured API Key
# -------------------------
@st.cache_resource
def load_hf_client():
    try:
        return InferenceClient(token=HF_API_KEY)
    except ImportError:
        st.error("huggingface_hub library not installed. Please install it with: pip install huggingface_hub")
        return None
    except Exception as e:
        st.error(f"Error creating client: {e}")
        return None

def map_sentiment_label(label, model_name):
    """Map model-specific labels to standard sentiment labels"""
    # For cardiffnlp/twitter-roberta-base-sentiment model
    if "twitter-roberta" in model_name:
        label_mapping = {
            "LABEL_0": "negative",
            "LABEL_1": "neutral", 
            "LABEL_2": "positive"
        }
        return label_mapping.get(label, label)
    
    # For distilbert sentiment model
    elif "distilbert" in model_name:
        label_mapping = {
            "LABEL_0": "negative",
            "LABEL_1": "positive"
        }
        return label_mapping.get(label, label)
    
    # For BERT multilingual sentiment model (5-star ratings)
    elif "bert-base-multilingual" in model_name:
        # Convert star rating to sentiment
        if "1 star" in label or "2 stars" in label:
            return "negative"
        elif "3 stars" in label:
            return "neutral"
        elif "4 stars" in label or "5 stars" in label:
            return "positive"
    
    # Default mapping for unknown models
    label_mapping = {
        "LABEL_0": "negative",
        "LABEL_1": "neutral", 
        "LABEL_2": "positive",
        "negative": "negative",
        "neutral": "neutral",
        "positive": "positive"
    }
    return label_mapping.get(label, label)

def predict_texts_hf(texts, client, model_name="cardiffnlp/twitter-roberta-base-sentiment"):
    """Call Hugging Face API for text classification with fallback."""
    results = []
    for txt in texts:
        try:
            # Use the text_classification method
            response = client.text_classification(txt, model=model_name)
            
            if not response:
                results.append({"label": "unknown", "confidence": 0.0, "distribution": {}})
                continue
                
            # Process the response
            distribution = {}
            for item in response:
                label = map_sentiment_label(item['label'], model_name)
                score = item['score']
                distribution[label] = score
            
            # Find the highest scoring sentiment
            top = max(response, key=lambda x: x["score"])
            sentiment_label = map_sentiment_label(top["label"], model_name)
            
            results.append({
                "label": sentiment_label,
                "confidence": float(top["score"]),
                "distribution": distribution
            })
                
        except Exception as e:
            st.error(f"Error from Hugging Face API: {e}")
            # Fallback to direct API call
            try:
                fallback_result = hugging_face_direct_api(txt, model_name)
                results.append(fallback_result)
            except:
                results.append({"label": "error", "confidence": 0.0, "distribution": {}})
    
    return results

def hugging_face_direct_api(text, model_name, max_retries=3):
    """Direct API call as fallback when InferenceClient fails."""
    api_url = f"https://api-inference.huggingface.co/models/{model_name}"
    headers = {"Authorization": f"Bearer {HF_API_KEY}"}
    
    for attempt in range(max_retries):
        try:
            response = requests.post(api_url, headers=headers, json={"inputs": text}, timeout=30)
            
            if response.status_code == 200:
                data = response.json()
                if isinstance(data, list) and len(data) > 0:
                    sentiment_data = data[0]
                    distribution = {}
                    for entry in sentiment_data:
                        label = map_sentiment_label(entry["label"], model_name)
                        score = float(entry["score"])
                        distribution[label] = score
                    
                    top = max(sentiment_data, key=lambda x: x["score"])
                    sentiment_label = map_sentiment_label(top["label"], model_name)
                    
                    return {
                        "label": sentiment_label,
                        "confidence": float(top["score"]),
                        "distribution": distribution
                    }
            
            # If model is loading, wait and retry
            elif response.status_code == 503:
                loading_time = response.json().get("estimated_time", 30)
                st.warning(f"Model is loading. Waiting {loading_time} seconds...")
                time.sleep(loading_time)
                continue
                
        except Exception as e:
            st.error(f"Direct API call failed: {e}")
            time.sleep(2)  # Wait before retry
    
    return {"label": "error", "confidence": 0.0, "distribution": {}}

# -------------------------
# Improved Keyword Extraction
# -------------------------
def extract_keywords(text, top_n=5):
    """Improved keyword extraction using NLP techniques"""
    if not text or not isinstance(text, str) or not text.strip():
        return []
    
    try:
        # Clean the text
        text = text.lower().strip()
        
        # Remove URLs
        text = re.sub(r'http\S+', '', text)
        
        # Remove punctuation
        text = text.translate(str.maketrans('', '', string.punctuation))
        
        # Tokenize
        words = word_tokenize(text)
        
        # Remove stopwords
        stop_words = set(stopwords.words('english'))
        words = [word for word in words if word not in stop_words and len(word) > 2]
        
        # Remove common non-meaningful words
        common_words = {
            'product', 'service', 'item', 'thing', 'stuff', 'would', 'could', 
            'should', 'really', 'very', 'much', 'many', 'also', 'like', 'get',
            'this', 'that', 'the', 'and', 'for', 'with', 'have', 'has', 'had',
            'was', 'were', 'are', 'is', 'be', 'been', 'being', 'all', 'just'
        }
        words = [word for word in words if word not in common_words]
        
        # Get word frequencies
        word_freq = Counter(words)
        
        # Return top N keywords
        return [word for word, count in word_freq.most_common(top_n)]
    
    except Exception as e:
        # Fallback to simple word extraction
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        stop_words = {'the', 'and', 'is', 'in', 'it', 'to', 'of', 'for', 'on', 'with', 'that', 'this', 'was', 'are', 'as', 'be', 'at', 'by', 'an', 'all', 'just'}
        filtered_words = [word for word in words if word not in stop_words]
        counter = Counter(filtered_words)
        return [word for word, count in counter.most_common(top_n)]

def extract_keywords_for_corpus(texts, top_n=5):
    """Extract keywords for multiple texts"""
    return [extract_keywords(text, top_n) for text in texts]

# -------------------------
# Demo Mode Functions
# -------------------------
def demo_sentiment_analysis(text):
    """Mock sentiment analysis for demo mode."""
    # Simple rule-based sentiment analysis
    positive_words = ["good", "great", 'excellent', 'amazing', 'wonderful', 'love', 'like', 'awesome', 'fantastic', 'best']
    negative_words = ["bad", 'terrible', 'awful', 'hate', 'dislike', 'worst', 'horrible', 'poor', 'boring', 'waste']
    
    text_lower = text.lower()
    positive_count = sum(1 for word in positive_words if word in text_lower)
    negative_count = sum(1 for word in negative_words if word in text_lower)
    
    if positive_count > negative_count:
        return {"label": "positive", "confidence": 0.7 + (positive_count * 0.05), "distribution": {"positive": 0.7, "neutral": 0.2, "negative": 0.1}}
    elif negative_count > positive_count:
        return {"label": "negative", "confidence": 0.7 + (negative_count * 0.05), "distribution": {"positive": 0.1, "neutral": 0.2, "negative": 0.7}}
    else:
        return {"label": "neutral", "confidence": 0.6, "distribution": {"positive": 0.3, "neutral": 0.4, "negative": 0.3}}

def demo_batch_analysis(texts):
    """Mock batch analysis for demo mode."""
    return [demo_sentiment_analysis(text) for text in texts]

# -------------------------
# Sidebar Settings
# -------------------------
st.sidebar.header("⚙️ Settings")

# Initialize and load the Hugging Face client
client = load_hf_client()
if client is not None:
    st.sidebar.success("✅ Hugging Face API connected successfully!")
else:
    st.sidebar.error("❌ Failed to connect to Hugging Face API")

# Model selection
model_name = st.sidebar.selectbox(
    "Choose a model",
    options=[
        "cardiffnlp/twitter-roberta-base-sentiment",  # 3 classes
        "distilbert-base-uncased-finetuned-sst-2-english",  # 2 classes
        "nlptown/bert-base-multilingual-uncased-sentiment"  # 5 classes
    ],
    index=0
)

top_n_keywords = st.sidebar.number_input("Top N keywords per text", min_value=1, max_value=10, value=5)

# Demo mode toggle (disabled since we have a real API key)
demo_mode = st.sidebar.checkbox("Enable Demo Mode (no API needed)", value=False)
if demo_mode:
    st.sidebar.info("Demo mode enabled. Using mock data for demonstration.")

# -------------------------
# Main UI
# -------------------------
# Header with gradient background
st.markdown("""
<div class="header">
    <h1>Sentiment Analyzer</h1>
    <h2>AI in Action - Understand the emotions behind your text</h2>
</div>
""", unsafe_allow_html=True)

tabs = st.tabs(["📝 Single Text", "📂 Batch Upload", "📊 Compare Datasets", "📈 Accuracy Report", "💾 Export Results"])

# -------------------------
# 1. Single Text
# -------------------------
with tabs[0]:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Analyze a single text")
    text = st.text_area("Enter text to analyze", height=150, placeholder="Type your text here...", 
                       value="", help="Enter any text to analyze its sentiment")
    
    if st.button("Analyze text", key="analyze_single"):
        if not text.strip():
            st.error("Please enter some text to analyze.")
        else:
            with st.spinner("Analyzing sentiment..."):
                if demo_mode:
                    preds = [demo_sentiment_analysis(text)]
                    st.info("Demo mode: Using mock sentiment analysis")
                else:
                    preds = predict_texts_hf([text], client, model_name=model_name)
                
                keywords = extract_keywords(text, top_n=top_n_keywords)
                res = preds[0]
                
                st.markdown('<div class="success-msg">Analysis Complete!</div>', unsafe_allow_html=True)
                
                col1, col2 = st.columns(2)
                with col1:
                    # Color code based on sentiment
                    sentiment_class = "positive" if res["label"] == "positive" else "negative" if res["label"] == "negative" else "neutral"
                    st.markdown(f"<h2 class='{sentiment_class}'>Sentiment: {res['label'].capitalize()}</h2>", unsafe_allow_html=True)
                    st.write(f"**Confidence:** {res['confidence']:.3f}")
                    
                    # Star rating based on sentiment
                    if res["label"] == "positive":
                        st.write("**Star Rating:** ⭐⭐⭐⭐⭐")
                    elif res["label"] == "negative":
                        st.write("**Star Rating:** ⭐")
                    else:
                        st.write("**Star Rating:** ⭐⭐⭐")
                        
                with col2:
                    st.write("**Keywords:**")
                    if keywords:
                        for i, keyword in enumerate(keywords, 1):
                            st.write(f"{i}. {keyword}")
                    else:
                        st.write("No significant keywords found")
                
                st.write("**Probability Distribution:**")
                for label, score in res["distribution"].items():
                    # Create a progress bar for each sentiment
                    label_class = "positive" if label == "positive" else "negative" if label == "negative" else "neutral"
                    st.markdown(f"<span class='{label_class}'>{label.capitalize()}:</span> {score:.3f}", unsafe_allow_html=True)
                    st.progress(score)
                    
                # Show the original text for context
                st.write("**Original Text:**")
                st.info(text)
    st.markdown('</div>', unsafe_allow_html=True)

# -------------------------
# 2. Batch Upload
# -------------------------
with tabs[1]:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Batch upload (CSV)")
    uploaded_file = st.file_uploader("Upload CSV with a text column", type=["csv"], help="Upload a CSV file containing text to analyze")
    
    if uploaded_file is not None:
        df = read_csv_with_encoding(uploaded_file)
        if df is not None:
            st.write("Preview of uploaded data:")
            st.dataframe(df.head())
            
            text_column = st.selectbox("Select text column", options=list(df.columns))
            
            if st.button("Run batch analysis", key="analyze_batch"):
                with st.spinner("Analyzing batch data..."):
                    texts = df[text_column].astype(str).fillna("").tolist()
                    # Limit to first 50 rows for demo purposes
                    if len(texts) > 50:
                        st.info(f"Analyzing first 50 of {len(texts)} rows. For full analysis, consider using a paid API plan.")
                        texts = texts[:50]
                    
                    if demo_mode:
                        preds = demo_batch_analysis(texts)
                        st.info("Demo mode: Using mock sentiment analysis")
                    else:
                        preds = predict_texts_hf(texts, client, model_name=model_name)
                    
                    keywords_list = extract_keywords_for_corpus(texts, top_n=top_n_keywords)
                    df_results = df.iloc[:len(preds)].copy()
                    df_results["sentiment"] = [p["label"] for p in preds]
                    df_results["confidence"] = [p["confidence"] for p in preds]
                    df_results["keywords"] = [", ".join(k) for k in keywords_list]
                    
                    # Store results in session state for export
                    st.session_state.df_results = df_results
                    
                    st.markdown('<div class="success-msg">Batch Analysis Complete!</div>', unsafe_allow_html=True)
                    
                    st.subheader("Analysis Results")
                    st.dataframe(df_results)
                    
                    st.subheader("Sentiment Distribution")
                    sentiment_counts = df_results["sentiment"].value_counts()
                    
                    # Create a proper visualization
                    fig, ax = plt.subplots(figsize=(8, 6))
                    colors = ['#dc3545' if s == 'negative' else '#ffc107' if s == 'neutral' else '#28a745' for s in sentiment_counts.index]
                    bars = ax.bar(sentiment_counts.index, sentiment_counts.values, color=colors)
                    
                    # Add value labels on bars
                    for bar in bars:
                        height = bar.get_height()
                        ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                                f'{int(height)}', ha='center', va='bottom')
                    
                    ax.set_ylabel('Count')
                    ax.set_title('Sentiment Distribution')
                    plt.xticks(rotation=45)
                    st.pyplot(fig)
                    
                    # Add download button
                    csv = df_results.to_csv(index=False)
                    st.download_button(
                        label="Download results as CSV",
                        data=csv,
                        file_name="sentiment_analysis_results.csv",
                        mime="text/csv"
                    )
    st.markdown('</div>', unsafe_allow_html=True)

# -------------------------
# 3. Compare Datasets
# -------------------------
with tabs[2]:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Compare two datasets")
    file_a = st.file_uploader("Upload dataset A", type=["csv"], key="a")
    file_b = st.file_uploader("Upload dataset B", type=["csv"], key="b")
    
    # Visualization type selection
    graph_type = st.selectbox(
        "Select visualization type",
        options=["Bar Chart", "Pie Chart", "Histogram", "Line Chart"],
        index=0
    )
    
    if file_a and file_b:
        df_a = read_csv_with_encoding(file_a)
        df_b = read_csv_with_encoding(file_b)
        
        if df_a is not None and df_b is not None:
            col_a = st.selectbox("Text column in A", options=list(df_a.columns), key="ca")
            col_b = st.selectbox("Text column in B", options=list(df_b.columns), key="cb")
            
            if st.button("Compare", key="compare_datasets"):
                with st.spinner("Analyzing datasets..."):
                    # Sample data for comparison (limit to 30 each)
                    sample_a = df_a[col_a].astype(str).fillna("").tolist()[:30]
                    sample_b = df_b[col_b].astype(str).fillna("").tolist()[:30]
                    
                    if demo_mode:
                        preds_a = demo_batch_analysis(sample_a)
                        preds_b = demo_batch_analysis(sample_b)
                        st.info("Demo mode: Using mock sentiment analysis")
                    else:
                        preds_a = predict_texts_hf(sample_a, client, model_name=model_name)
                        preds_b = predict_texts_hf(sample_b, client, model_name=model_name)
                    
                    counts_a = pd.Series([p["label"] for p in preds_a]).value_counts()
                    counts_b = pd.Series([p["label"] for p in preds_b]).value_counts()
                    
                    comp = pd.DataFrame({"Dataset A": counts_a, "Dataset B": counts_b}).fillna(0)
                    
                    # Store comparison results for export and accuracy report
                    st.session_state.comparison_results = {
                        "dataset_a": sample_a,
                        "dataset_b": sample_b,
                        "preds_a": preds_a,
                        "preds_b": preds_b,
                        "comparison_df": comp
                    }
                    
                    st.markdown('<div class="success-msg">Comparison Complete!</div>', unsafe_allow_html=True)
                    
                    st.subheader("Comparison Results")
                    st.dataframe(comp)
                    
                    # Create visualization based on user selection
                    fig, ax = plt.subplots(figsize=(10, 6))
                    
                    if graph_type == "Bar Chart":
                        # Set up the bar positions
                        x = np.arange(len(comp.index))
                        width = 0.35
                        
                        # Create bars for each dataset
                        bars1 = ax.bar(x - width/2, comp["Dataset A"], width, label='Dataset A', alpha=0.8, color='#6e8efb')
                        bars2 = ax.bar(x + width/2, comp["Dataset B"], width, label='Dataset B', alpha=0.8, color='#a777e3')
                        
                        # Add value labels on bars
                        for bars in [bars1, bars2]:
                            for bar in bars:
                                height = bar.get_height()
                                if height > 0:
                                    ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                                            f'{int(height)}', ha='center', va='bottom')
                        
                        ax.set_xlabel('Sentiment')
                        ax.set_ylabel('Count')
                        ax.set_title('Sentiment Comparison Between Datasets')
                        ax.set_xticks(x)
                        ax.set_xticklabels(comp.index)
                        ax.legend()
                    
                    elif graph_type == "Pie Chart":
                        # Create subplots for each dataset
                        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 6))
                        
                        # Dataset A pie chart
                        colors = ['#dc3545', '#ffc107', '#28a745'][:len(comp.index)]
                        ax1.pie(comp["Dataset A"], labels=comp.index, autopct='%1.1f%%', startangle=90, colors=colors)
                        ax1.set_title('Dataset A Sentiment Distribution')
                        
                        # Dataset B pie chart
                        ax2.pie(comp["Dataset B"], labels=comp.index, autopct='%1.1f%%', startangle=90, colors=colors)
                        ax2.set_title('Dataset B Sentiment Distribution')
                        
                        # Equal aspect ratio ensures that pie is drawn as a circle
                        ax1.axis('equal')
                        ax2.axis('equal')
                    
                    elif graph_type == "Histogram":
                        # Prepare data for histogram
                        sentiment_labels = comp.index.tolist()
                        a_values = comp["Dataset A"].tolist()
                        b_values = comp["Dataset B"].tolist()
                        
                        x = np.arange(len(sentiment_labels))
                        width = 0.35
                        
                        ax.bar(x - width/2, a_values, width, label='Dataset A', alpha=0.8, color='#6e8efb')
                        ax.bar(x + width/2, b_values, width, label='Dataset B', alpha=0.8, color='#a777e3')
                        
                        ax.set_xlabel('Sentiment')
                        ax.set_ylabel('Count')
                        ax.set_title('Sentiment Comparison Histogram')
                        ax.set_xticks(x)
                        ax.set_xticklabels(sentiment_labels)
                        ax.legend()
                    
                    elif graph_type == "Line Chart":
                        # Prepare data for line chart
                        sentiment_labels = comp.index.tolist()
                        a_values = comp["Dataset A"].tolist()
                        b_values = comp["Dataset B"].tolist()
                        
                        x = range(len(sentiment_labels))
                        
                        ax.plot(x, a_values, marker='o', label='Dataset A', linewidth=2, color='#6e8efb')
                        ax.plot(x, b_values, marker='s', label='Dataset B', linewidth=2, color='#a777e3')
                        
                        ax.set_xlabel('Sentiment')
                        ax.set_ylabel('Count')
                        ax.set_title('Sentiment Comparison Line Chart')
                        ax.set_xticks(x)
                        ax.set_xticklabels(sentiment_labels)
                        ax.legend()
                        ax.grid(True, linestyle='--', alpha=0.7)
                    
                    st.pyplot(fig)
                    
                    # Add download button for comparison results
                    comp_csv = comp.to_csv()
                    st.download_button(
                        label="Download comparison results",
                        data=comp_csv,
                        file_name="dataset_comparison.csv",
                        mime="text/csv",
                        help="Download the comparison results for use in accuracy reports"
                    )
                    
                    # Create detailed comparison results for download
                    detailed_comparison = pd.DataFrame({
                        "Dataset A Text": sample_a,
                        "Dataset A Sentiment": [p["label"] for p in preds_a],
                        "Dataset A Confidence": [p["confidence"] for p in preds_a],
                        "Dataset B Text": sample_b,
                        "Dataset B Sentiment": [p["label"] for p in preds_b],
                        "Dataset B Confidence": [p["confidence"] for p in preds_b]
                    })
                    
                    detailed_csv = detailed_comparison.to_csv(index=False)
                    st.download_button(
                        label="Download detailed comparison results",
                        data=detailed_csv,
                        file_name="detailed_dataset_comparison.csv",
                        mime="text/csv",
                        help="Download detailed comparison results with text and sentiment scores"
                    )
    st.markdown('</div>', unsafe_allow_html=True)

# -------------------------
# 4. Accuracy Report
# -------------------------
with tabs[3]:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Accuracy report")
    
    use_comparison_data = st.checkbox("Use comparison data for accuracy report", value=False)
    
    if use_comparison_data and st.session_state.comparison_results is not None:
        st.info("Using comparison data for accuracy report")
        
        comparison_data = st.session_state.comparison_results
        preds_a = comparison_data["preds_a"]
        preds_b = comparison_data["preds_b"]
        
        # Get gold and predicted labels
        gold_labels = [p["label"] for p in preds_a]
        pred_labels = [p["label"] for p in preds_b]
        
        # Only compare where both are not "unknown" or "error"
        valid_indices = [
            i for i, (g, p) in enumerate(zip(gold_labels, pred_labels))
            if g in ["positive", "neutral", "negative"] and p in ["positive", "neutral", "negative"]
        ]
        gold_labels_valid = [gold_labels[i] for i in valid_indices]
        pred_labels_valid = [pred_labels[i] for i in valid_indices]
        
        if len(gold_labels_valid) == 0 or len(pred_labels_valid) == 0:
            st.error("No valid comparison data available for accuracy calculation.")
        else:
            st.subheader("Classification Report")
            report = classification_report(gold_labels_valid, pred_labels_valid, output_dict=True, zero_division=0)
            st.dataframe(pd.DataFrame(report).transpose())
            
            st.subheader("Confusion Matrix")
            labels = sorted(list(set(gold_labels_valid + pred_labels_valid)))
            cm = confusion_matrix(gold_labels_valid, pred_labels_valid, labels=labels)
            
            fig, ax = plt.subplots(figsize=(8, 6))
            im = ax.imshow(cm, cmap="Blues")
            ax.set_xticks(range(len(labels)))
            ax.set_xticklabels(labels, rotation=45)
            ax.set_yticks(range(len(labels)))
            ax.set_yticklabels(labels)
            ax.set_xlabel("Predicted (Dataset B)")
            ax.set_ylabel("Actual (Dataset A)")
            ax.set_title("Confusion Matrix: Dataset A vs Dataset B")
            
            for i in range(len(labels)):
                for j in range(len(labels)):
                    ax.text(j, i, cm[i, j], ha="center", va="center", 
                           color="white" if cm[i, j] > cm.max()/2 else "black")
            
            plt.colorbar(im)
            st.pyplot(fig)
            
            total = np.sum(cm)
            correct = np.trace(cm)
            if total > 0:
                accuracy = correct / total
                st.metric("Overall Accuracy", f"{accuracy:.2%}")
            else:
                st.metric("Overall Accuracy", "N/A")
    else:
        # Traditional accuracy report with labeled data
        labeled = st.file_uploader("Upload CSV with text and label columns", type=["csv"])
        
        if labeled is not None:
            df_lab = read_csv_with_encoding(labeled)
            if df_lab is not None:
                # Check if the required columns exist
                text_cols = [col for col in df_lab.columns if 'text' in col.lower()]
                label_cols = [col for col in df_lab.columns if 'label' in col.lower()]
                
                if text_cols and label_cols:
                    text_col = st.selectbox("Select text column", options=text_cols)
                    label_col = st.selectbox("Select label column", options=label_cols)
                    
                    st.write("Preview of labeled data:")
                    st.dataframe(df_lab.head())
                    
                    if st.button("Generate Accuracy Report", key="accuracy_report"):
                        with st.spinner("Generating accuracy report..."):
                            # Sample data for analysis
                            sample_texts = df_lab[text_col].astype(str).fillna("").tolist()[:30]
                            true_labels = df_lab[label_col].astype(str).fillna("").tolist()[:30]
                            
                            if demo_mode:
                                preds = demo_batch_analysis(sample_texts)
                                st.info("Demo mode: Using mock sentiment analysis")
                            else:
                                preds = predict_texts_hf(sample_texts, client, model_name=model_name)
                            
                            pred_labels = [p["label"] for p in preds]
                            
                            # Limit to the same number of samples
                            min_len = min(len(true_labels), len(pred_labels))
                            true_labels = true_labels[:min_len]
                            pred_labels = pred_labels[:min_len]
                            
                            st.markdown('<div class="success-msg">Accuracy Report Generated!</div>', unsafe_allow_html=True)
                            
                            st.subheader("Classification Report")
                            report = classification_report(true_labels, pred_labels, output_dict=True, zero_division=0)
                            st.dataframe(pd.DataFrame(report).transpose())
                            
                            st.subheader("Confusion Matrix")
                            labels = sorted(list(set(true_labels + pred_labels)))
                            cm = confusion_matrix(true_labels, pred_labels, labels=labels)
                            
                            fig, ax = plt.subplots(figsize=(8, 6))
                            im = ax.imshow(cm, cmap="Blues")
                            ax.set_xticks(range(len(labels)))
                            ax.set_xticklabels(labels, rotation=45)
                            ax.set_yticks(range(len(labels)))
                            ax.set_yticklabels(labels)
                            ax.set_xlabel("Predicted")
                            ax.set_ylabel("Actual")
                            ax.set_title("Confusion Matrix")
                            
                            # Add text annotations
                            for i in range(len(labels)):
                                for j in range(len(labels)):
                                    ax.text(j, i, cm[i, j], ha="center", va="center", 
                                           color="white" if cm[i, j] > cm.max()/2 else "black")
                            
                            plt.colorbar(im)
                            st.pyplot(fig)
                            
                            # Calculate accuracy metrics
                            accuracy = np.trace(cm) / np.sum(cm)
                            st.metric("Overall Accuracy", f"{accuracy:.2%}")
                else:
                    st.error("Uploaded file must contain columns with 'text' and 'label' in their names")
    st.markdown('</div>', unsafe_allow_html=True)

# -------------------------
# 5. Export Results
# -------------------------
with tabs[4]:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Export Results")
    
    # Collect all history/results from session state
    history = {
        "batch_results": st.session_state.get("df_results"),
        "comparison_results": st.session_state.get("comparison_results"),
    }
    
    # Helper to export graphs as images
    def export_graph(fig, filename):
        buf = io.BytesIO()
        fig.savefig(buf, format="png")
        buf.seek(0)
        return buf

    # Show batch results
    if history["batch_results"] is not None:
        st.markdown('<div class="success-msg">Analysis results available for export!</div>', unsafe_allow_html=True)
        st.dataframe(history["batch_results"].head())
        cols = st.columns(6)
        # CSV Export
        csv = history["batch_results"].to_csv(index=False)
        cols[0].download_button(
            label="📄 CSV",
            data=csv,
            file_name="sentiment_analysis.csv",
            mime="text/csv",
            help="Download as CSV file"
        )
        # Excel Export
        excel_buffer = create_excel_export(history["batch_results"])
        if excel_buffer:
            cols[1].download_button(
                label="📊 Excel",
                data=excel_buffer,
                file_name="sentiment_analysis.xlsx",
                mime="application/vnd.ms-excel",
                help="Download as Excel file"
            )
        # HTML Export
        html = create_html_export(history["batch_results"])
        cols[2].download_button(
            label="🌐 HTML",
            data=html,
            file_name="sentiment_analysis.html",
            mime="text/html",
            help="Download as HTML file"
        )
        # Word Export
        word_buffer = create_word_export(history["batch_results"])
        if word_buffer:
            cols[3].download_button(
                label="📝 Word",
                data=word_buffer,
                file_name="sentiment_analysis.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                help="Download as Word document"
            )
        # PDF Export
        pdf_buffer = create_pdf_export(history["batch_results"])
        if pdf_buffer:
            cols[4].download_button(
                label="📑 PDF",
                data=pdf_buffer,
                file_name="sentiment_analysis.pdf",
                mime="application/pdf",
                help="Download as PDF document"
            )
        # JSON Export
        json_data = create_json_export(history["batch_results"])
        cols[5].download_button(
            label="🟧 JSON",
            data=json_data,
            file_name="sentiment_analysis.json",
            mime="application/json",
            help="Download as JSON file"
        )
        # Export batch graph
        st.subheader("Download Sentiment Distribution Graph")
        sentiment_counts = history["batch_results"]["sentiment"].value_counts()
        fig, ax = plt.subplots(figsize=(8, 6))
        colors = ['#dc3545' if s == 'negative' else '#ffc107' if s == 'neutral' else '#28a745' for s in sentiment_counts.index]
        bars = ax.bar(sentiment_counts.index, sentiment_counts.values, color=colors)
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 0.1, f'{int(height)}', ha='center', va='bottom')
        ax.set_ylabel('Count')
        ax.set_title('Sentiment Distribution')
        plt.xticks(rotation=45)
        graph_buf = export_graph(fig, "sentiment_distribution.png")
        st.download_button(
            label="Download Sentiment Distribution Graph (PNG)",
            data=graph_buf,
            file_name="sentiment_distribution.png",
            mime="image/png"
        )
        plt.close(fig)
    
    # Show comparison results
    if history["comparison_results"] is not None:
        st.subheader("Download Comparison Results and Graphs")
        comp_df = history["comparison_results"]["comparison_df"]
        st.dataframe(comp_df)
        comp_csv = comp_df.to_csv()
        st.download_button(
            label="Download comparison results (CSV)",
            data=comp_csv,
            file_name="dataset_comparison.csv",
            mime="text/csv"
        )
        # Export comparison bar chart
        fig, ax = plt.subplots(figsize=(10, 6))
        x = np.arange(len(comp_df.index))
        width = 0.35
        bars1 = ax.bar(x - width/2, comp_df["Dataset A"], width, label='Dataset A', alpha=0.8, color='#6e8efb')
        bars2 = ax.bar(x + width/2, comp_df["Dataset B"], width, label='Dataset B', alpha=0.8, color='#a777e3')
        for bars in [bars1, bars2]:
            for bar in bars:
                height = bar.get_height()
                if height > 0:
                    ax.text(bar.get_x() + bar.get_width()/2., height + 0.1, f'{int(height)}', ha='center', va='bottom')
        ax.set_xlabel('Sentiment')
        ax.set_ylabel('Count')
        ax.set_title('Sentiment Comparison Between Datasets')
        ax.set_xticks(x)
        ax.set_xticklabels(comp_df.index)
        ax.legend()
        comp_graph_buf = export_graph(fig, "comparison_bar_chart.png")
        st.download_button(
            label="Download Comparison Bar Chart (PNG)",
            data=comp_graph_buf,
            file_name="comparison_bar_chart.png",
            mime="image/png"
        )
        plt.close(fig)
        # Export detailed comparison
        detailed_comparison = pd.DataFrame({
            "Dataset A Text": history["comparison_results"]["dataset_a"],
            "Dataset A Sentiment": [p["label"] for p in history["comparison_results"]["preds_a"]],
            "Dataset A Confidence": [p["confidence"] for p in history["comparison_results"]["preds_a"]],
            "Dataset B Text": history["comparison_results"]["dataset_b"],
            "Dataset B Sentiment": [p["label"] for p in history["comparison_results"]["preds_b"]],
            "Dataset B Confidence": [p["confidence"] for p in history["comparison_results"]["preds_b"]]
        })
        detailed_csv = detailed_comparison.to_csv(index=False)
        st.download_button(
            label="Download detailed comparison results (CSV)",
            data=detailed_csv,
            file_name="detailed_dataset_comparison.csv",
            mime="text/csv"
        )
    # If no results, show sample data
    if history["batch_results"] is None and history["comparison_results"] is None:
        st.markdown('<div class="info-msg">No analysis results available for export. Please run a batch analysis or comparison first.</div>', unsafe_allow_html=True)
        # Sample data for demonstration
        sample_data = {
            'text': ['I love this product!', 'This is terrible', 'It is okay'],
            'sentiment': ['positive', 'negative', 'neutral'],
            'confidence': [0.95, 0.87, 0.65],
            'keywords': ['love, product', 'terrible', 'okay']
        }
        sample_df = pd.DataFrame(sample_data)
        
        st.write("Sample export format:")
        st.dataframe(sample_df)
        
        # Export sample data
        st.subheader("Try export with sample data:")
        cols = st.columns(6)
        # CSV Export
        csv = sample_df.to_csv(index=False)
        cols[0].download_button(
            label="📄 CSV",
            data=csv,
            file_name="sample_sentiment.csv",
            mime="text/csv"
        )
        # Excel Export
        excel_buffer = create_excel_export(sample_df)
        if excel_buffer:
            cols[1].download_button(
                label="📊 Excel",
                data=excel_buffer,
                file_name="sample_sentiment.xlsx",
                mime="application/vnd.ms-excel"
            )
        # HTML Export
        html = create_html_export(sample_df)
        cols[2].download_button(
            label="🌐 HTML",
            data=html,
            file_name="sample_sentiment.html",
            mime="text/html"
        )
        # Word Export
        word_buffer = create_word_export(sample_df)
        if word_buffer:
            cols[3].download_button(
                label="📝 Word",
                data=word_buffer,
                file_name="sample_sentiment.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        # PDF Export
        pdf_buffer = create_pdf_export(sample_df)
        if pdf_buffer:
            cols[4].download_button(
                label="📑 PDF",
                data=pdf_buffer,
                file_name="sample_sentiment.pdf",
                mime="application/pdf"
            )
        # JSON Export
        json_data = create_json_export(sample_df)
        cols[5].download_button(
            label="🟧 JSON",
            data=json_data,
            file_name="sample_sentiment.json",
            mime="application/json"
        )
    st.markdown('</div>', unsafe_allow_html=True)
    

    # --- Export Single Text Results ---
    if "analyze_single" in st.session_state and st.session_state.get("single_result"):
        single_res = st.session_state["single_result"]
        single_df = pd.DataFrame([single_res])
        st.subheader("Download Single Text Result")
        st.dataframe(single_df)
        st.download_button(
            label="Download single text result (CSV)",
            data=single_df.to_csv(index=False),
            file_name="single_text_result.csv",
            mime="text/csv"
        )
        st.download_button(
            label="Download single text result (JSON)",
            data=create_json_export(single_df),
            file_name="single_text_result.json",
            mime="application/json"
        )
